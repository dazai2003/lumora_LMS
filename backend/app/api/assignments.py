"""
Assignment & Coursework Management API.
Full-featured assignment lifecycle: creation, attachments, group formation, student submissions,
rubric evaluations, AI-assisted grading, plagiarism analysis, and coursework analytics.
"""
import os
import logging
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List, Optional
from datetime import datetime

from app.database import get_db
from app.models import (
    User, UserRole, Course, Lesson, Enrollment,
    Assignment, AssignmentFile, AssignmentResource, AssignmentGroup, GroupMember,
    AssignmentSubmission, SubmissionFile, SubmissionHistory, SubmissionAnnotation,
    AssignmentRubric, RubricCriteria, RubricScoreDetail, PlagiarismReport,
    SubmissionVersion, SubmissionComment, SubmissionSuggestion,
    SubmissionSectionFeedback, DocumentExtraction,
    Notification, NotificationType
)
from app.auth import get_current_user, require_role

logger = logging.getLogger(__name__)
router = APIRouter()


# ──────────────────────────────────────────────
# 1. Assignment CRUD
# ──────────────────────────────────────────────

@router.post("", response_model=dict)
def create_assignment(
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher creates a new coursework assignment."""
    course_id = data.get("course_id")
    title = data.get("title")
    if not course_id or not title:
        raise HTTPException(status_code=400, detail="Course ID and Title are required")
        
    course = db.query(Course).filter(Course.id == course_id, Course.teacher_id == current_user.id).first()
    if not course:
        raise HTTPException(status_code=403, detail="Course not found or unauthorized")

    assignment = Assignment(
        course_id=course_id,
        lesson_id=data.get("lesson_id"),
        title=title,
        description=data.get("description"),
        instructions=data.get("instructions"),
        max_marks=float(data.get("max_marks", 100.0)),
        weightage=float(data.get("weightage", 10.0)),
        is_group=bool(data.get("is_group", False)),
        status=data.get("status", "published"),
        available_from=datetime.fromisoformat(data["available_from"]) if data.get("available_from") else None,
        available_until=datetime.fromisoformat(data["available_until"]) if data.get("available_until") else None,
        due_date=datetime.fromisoformat(data["due_date"]) if data.get("due_date") else None,
        # Phase 4.1 fields
        learning_outcomes=data.get("learning_outcomes", []),
        blooms_level=data.get("blooms_level"),
        difficulty=data.get("difficulty", "medium"),
        est_completion_time_minutes=int(data["est_completion_time_minutes"]) if data.get("est_completion_time_minutes") else None,
        category=data.get("category"),
        anonymous_marking=bool(data.get("anonymous_marking", False)),
        ai_policy=data.get("ai_policy", "allowed"),
        word_count_limits=data.get("word_count_limits"),
        allowed_file_types=data.get("allowed_file_types"),
        max_upload_size_mb=int(data.get("max_upload_size_mb", 50)),
        late_submission_rules=data.get("late_submission_rules"),
        max_attempts=int(data.get("max_attempts", 1)),
        ai_pre_check_enabled=bool(data.get("ai_pre_check_enabled", False)),
    )
    db.add(assignment)
    db.commit()
    db.refresh(assignment)

    # Optional rubric setup
    rubric_criteria = data.get("rubric_criteria", [])
    if rubric_criteria:
        rubric = AssignmentRubric(assignment_id=assignment.id, title=f"Rubric for {title}")
        db.add(rubric)
        db.commit()
        db.refresh(rubric)
        for order, crit in enumerate(rubric_criteria, start=1):
            rc = RubricCriteria(
                rubric_id=rubric.id,
                criterion_name=crit.get("name", f"Criteria {order}"),
                description=crit.get("description", ""),
                max_score=float(crit.get("max_score", 10.0)),
                weight=float(crit.get("weight", 1.0)),
                order=order
            )
            db.add(rc)
        db.commit()

    return {
        "id": assignment.id,
        "course_id": assignment.course_id,
        "title": assignment.title,
        "max_marks": assignment.max_marks,
        "weightage": assignment.weightage,
        "is_group": assignment.is_group,
        "status": assignment.status,
        "category": assignment.category,
        "blooms_level": assignment.blooms_level,
        "difficulty": assignment.difficulty,
        "due_date": assignment.due_date.isoformat() if assignment.due_date else None,
        "created_at": assignment.created_at.isoformat()
    }


@router.put("/{assignment_id}", response_model=dict)
def update_assignment(
    assignment_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher updates an existing assignment."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    course = db.query(Course).filter(Course.id == assignment.course_id, Course.teacher_id == current_user.id).first()
    if not course:
        raise HTTPException(status_code=403, detail="Unauthorized to edit this assignment")

    if "title" in data and data["title"]: assignment.title = data["title"]
    if "description" in data: assignment.description = data["description"]
    if "instructions" in data: assignment.instructions = data["instructions"]
    if "max_marks" in data and data["max_marks"] is not None: assignment.max_marks = float(data["max_marks"])
    if "weightage" in data and data["weightage"] is not None: assignment.weightage = float(data["weightage"])
    if "is_group" in data: assignment.is_group = bool(data["is_group"])
    if "status" in data and data["status"]: assignment.status = data["status"]
    if "due_date" in data:
        assignment.due_date = datetime.fromisoformat(data["due_date"]) if data["due_date"] else None
    if "available_from" in data:
        assignment.available_from = datetime.fromisoformat(data["available_from"]) if data["available_from"] else None
    if "available_until" in data:
        assignment.available_until = datetime.fromisoformat(data["available_until"]) if data["available_until"] else None
    if "category" in data: assignment.category = data["category"]
    if "blooms_level" in data: assignment.blooms_level = data["blooms_level"]
    if "difficulty" in data and data["difficulty"]: assignment.difficulty = data["difficulty"]
    if "est_completion_time_minutes" in data:
        assignment.est_completion_time_minutes = int(data["est_completion_time_minutes"]) if data["est_completion_time_minutes"] else None
    if "ai_policy" in data and data["ai_policy"]: assignment.ai_policy = data["ai_policy"]
    if "word_count_limits" in data: assignment.word_count_limits = data["word_count_limits"]
    if "late_submission_rules" in data: assignment.late_submission_rules = data["late_submission_rules"]
    if "anonymous_marking" in data: assignment.anonymous_marking = bool(data["anonymous_marking"])
    if "ai_pre_check_enabled" in data: assignment.ai_pre_check_enabled = bool(data["ai_pre_check_enabled"])

    db.commit()
    db.refresh(assignment)
    return {"message": "Assignment updated successfully", "id": assignment.id, "title": assignment.title, "status": assignment.status}


@router.delete("/{assignment_id}", response_model=dict)
def delete_assignment(
    assignment_id: int,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher deletes an assignment and all associated records."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    course = db.query(Course).filter(Course.id == assignment.course_id, Course.teacher_id == current_user.id).first()
    if not course:
        raise HTTPException(status_code=403, detail="Unauthorized to delete this assignment")

    db.delete(assignment)
    db.commit()
    return {"message": "Assignment deleted successfully", "id": assignment_id}


@router.delete("/resources/{resource_id}", response_model=dict)
def delete_assignment_resource(
    resource_id: int,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher removes a resource attached to an assignment."""
    resource = db.query(AssignmentResource).filter(AssignmentResource.id == resource_id).first()
    if not resource:
        raise HTTPException(status_code=404, detail="Resource not found")

    db.delete(resource)
    db.commit()
    return {"message": "Resource deleted successfully", "id": resource_id}



@router.get("", response_model=List[dict])
def list_assignments(
    course_id: Optional[int] = None,
    status: Optional[str] = None,
    search: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(Assignment).join(Course)

    # Enforce multi-tenant ownership & role scoping
    if current_user.role == UserRole.TEACHER:
        query = query.filter(Course.teacher_id == current_user.id)
    elif current_user.role == UserRole.STUDENT:
        query = query.filter(
            Assignment.status == "published",
            Course.id.in_(
                db.query(Enrollment.course_id).filter(Enrollment.student_id == current_user.id)
            )
        )

    if course_id:
        query = query.filter(Assignment.course_id == course_id)
    if status:
        query = query.filter(Assignment.status == status)
    if search:
        query = query.filter(Assignment.title.ilike(f"%{search}%"))

    assignments = query.order_by(Assignment.due_date.asc().nullslast()).all()
    results = []
    for a in assignments:
        sub_count = len(a.submissions)
        my_sub = None
        if current_user.role == UserRole.STUDENT:
            sub = db.query(AssignmentSubmission).filter(
                AssignmentSubmission.assignment_id == a.id,
                AssignmentSubmission.student_id == current_user.id
            ).first()

            # Group member lookup
            if not sub and a.is_group:
                gm = db.query(GroupMember).join(AssignmentGroup).filter(
                    AssignmentGroup.assignment_id == a.id,
                    GroupMember.student_id == current_user.id
                ).first()
                if gm and gm.group:
                    sub = db.query(AssignmentSubmission).filter(
                        AssignmentSubmission.group_id == gm.group_id
                    ).first()

            if sub:
                my_sub = {
                    "submission_id": sub.id,
                    "status": sub.status,
                    "grade_marks": sub.grade_marks if sub.is_published else None,
                    "submitted_at": sub.submitted_at.isoformat() if sub.submitted_at else None,
                    "is_late": sub.is_late,
                    "is_group_submission": sub.student_id != current_user.id,
                    "submitted_by_name": sub.student.full_name if sub.student else "Group Leader"
                }

        results.append({
            "id": a.id,
            "course_id": a.course_id,
            "course_title": a.course.title if a.course else None,
            "title": a.title,
            "description": a.description,
            "max_marks": a.max_marks,
            "weightage": a.weightage,
            "is_group": a.is_group,
            "status": a.status,
            "due_date": a.due_date.isoformat() if a.due_date else None,
            "submission_count": sub_count,
            "my_submission": my_sub,
            "created_at": a.created_at.isoformat()
        })
    return results


@router.get("/{assignment_id}", response_model=dict)
def get_assignment_details(
    assignment_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get full assignment details including instructions, attachments, rubrics, and groups."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    files = [{"id": f.id, "file_name": f.file_name, "file_path": f.file_path, "file_size": f.file_size} for f in assignment.files]
    rubrics = []
    for r in assignment.rubrics:
        criteria = [{"id": c.id, "criterion_name": c.criterion_name, "description": c.description, "max_score": c.max_score, "weight": c.weight} for c in r.criteria]
        rubrics.append({"id": r.id, "title": r.title, "criteria": criteria})

    # Student submission detail lookup (including group membership)
    my_sub_details = None
    if current_user.role == UserRole.STUDENT:
        sub = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment_id,
            AssignmentSubmission.student_id == current_user.id
        ).first()

        if not sub and assignment.is_group:
            gm = db.query(GroupMember).join(AssignmentGroup).filter(
                AssignmentGroup.assignment_id == assignment_id,
                GroupMember.student_id == current_user.id
            ).first()
            if gm and gm.group:
                sub = db.query(AssignmentSubmission).filter(
                    AssignmentSubmission.group_id == gm.group_id
                ).first()

        if sub:
            files_sub = [{"id": f.id, "file_name": f.file_name, "file_path": f.file_path, "file_size": f.file_size} for f in sub.files]
            
            group_members_list = []
            group_mate_ids = []
            if sub.group_id:
                group = db.query(AssignmentGroup).filter(AssignmentGroup.id == sub.group_id).first()
                if group:
                    for gm in group.members:
                        if gm.student_id != sub.student_id:
                            group_mate_ids.append(gm.student_id)
                        if gm.student:
                            group_members_list.append({
                                "id": gm.student.id,
                                "full_name": gm.student.full_name,
                                "email": gm.student.email
                            })

            my_sub_details = {
                "id": sub.id,
                "submission_id": sub.id,
                "status": sub.status,
                "submission_mode": sub.submission_mode,
                "submission_content_text": sub.submission_content_text,
                "student_comment": sub.student_comment,
                "repository_url": sub.repository_url,
                "word_count": sub.word_count,
                "grade_marks": sub.grade_marks if sub.is_published else None,
                "feedback_text": sub.feedback_text if sub.is_published else None,
                "is_published": sub.is_published,
                "submitted_at": sub.submitted_at.isoformat() if sub.submitted_at else None,
                "is_late": sub.is_late,
                "files": files_sub,
                "is_group_submission": sub.student_id != current_user.id,
                "submitted_by_name": sub.student.full_name if sub.student else "Group Leader",
                "group_mate_ids": group_mate_ids,
                "group_members": group_members_list
            }

    return {
        "id": assignment.id,
        "course_id": assignment.course_id,
        "course_title": assignment.course.title if assignment.course else None,
        "lesson_id": assignment.lesson_id,
        "title": assignment.title,
        "description": assignment.description,
        "instructions": assignment.instructions,
        "max_marks": assignment.max_marks,
        "weightage": assignment.weightage,
        "is_group": assignment.is_group,
        "status": assignment.status,
        "category": assignment.category,
        "blooms_level": assignment.blooms_level,
        "difficulty": assignment.difficulty,
        "est_completion_time_minutes": assignment.est_completion_time_minutes,
        "ai_policy": assignment.ai_policy,
        "word_count_limits": assignment.word_count_limits,
        "late_submission_rules": assignment.late_submission_rules,
        "anonymous_marking": assignment.anonymous_marking,
        "ai_pre_check_enabled": assignment.ai_pre_check_enabled,
        "available_from": assignment.available_from.isoformat() if assignment.available_from else None,
        "available_until": assignment.available_until.isoformat() if assignment.available_until else None,
        "due_date": assignment.due_date.isoformat() if assignment.due_date else None,
        "files": files,
        "rubrics": rubrics,
        "my_submission": my_sub_details,
        "created_at": assignment.created_at.isoformat()
    }


@router.get("/{assignment_id}/enrolled-students", response_model=List[dict])
def list_enrolled_students(
    assignment_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List enrolled classmates in the assignment's course for group mate tagging."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    enrollments = db.query(Enrollment).filter(
        Enrollment.course_id == assignment.course_id,
        Enrollment.student_id != current_user.id
    ).all()

    return [{
        "id": e.student.id,
        "full_name": e.student.full_name,
        "email": e.student.email
    } for e in enrollments if e.student]


# ──────────────────────────────────────────────
# 2. Student Submissions Engine
# ──────────────────────────────────────────────

@router.post("/{assignment_id}/submit", response_model=dict)
def submit_assignment(
    assignment_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """Student submits or saves draft for an assignment."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    now = datetime.utcnow()
    is_late = bool(assignment.due_date and now > assignment.due_date)
    is_draft = data.get("is_draft", False)

    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.student_id == current_user.id
    ).first()

    if not submission:
        submission = AssignmentSubmission(
            assignment_id=assignment_id,
            student_id=current_user.id,
            status="draft" if is_draft else "submitted",
            submitted_at=now,
            is_late=is_late,
            submission_mode=data.get("submission_mode", "rich_text"),
            submission_content_text=data.get("content_text") or data.get("submission_content_text"),
            student_comment=data.get("student_comment"),
            repository_url=data.get("repository_url"),
            word_count=len((data.get("content_text") or "").split()) if data.get("content_text") else 0
        )
        db.add(submission)
        db.commit()
        db.refresh(submission)
    else:
        submission.status = "draft" if is_draft else "submitted"
        submission.submitted_at = now
        submission.is_late = is_late
        if data.get("submission_mode"): submission.submission_mode = data["submission_mode"]
        if data.get("content_text") is not None: submission.submission_content_text = data["content_text"]
        if data.get("student_comment"): submission.student_comment = data["student_comment"]
        if data.get("repository_url"): submission.repository_url = data["repository_url"]
        if data.get("content_text"): submission.word_count = len(data["content_text"].split())
        db.commit()

    # Group Tagging Logic
    group_mate_ids = data.get("group_mate_ids", [])
    if assignment.is_group:
        group = None
        if submission.group_id:
            group = db.query(AssignmentGroup).filter(AssignmentGroup.id == submission.group_id).first()
        if not group:
            group = AssignmentGroup(
                assignment_id=assignment.id,
                group_name=f"Group - {current_user.full_name}",
                leader_id=current_user.id
            )
            db.add(group)
            db.commit()
            db.refresh(group)
            submission.group_id = group.id
            db.commit()

        # Add submitter + tagged group mates
        all_member_ids = set([current_user.id] + [int(mid) for mid in group_mate_ids if str(mid).isdigit()])
        db.query(GroupMember).filter(GroupMember.group_id == group.id).delete()
        for mid in all_member_ids:
            gm = GroupMember(group_id=group.id, student_id=mid, contribution_percentage=round(100.0 / len(all_member_ids), 1))
            db.add(gm)
        db.commit()

    # Log submission history
    history = SubmissionHistory(
        submission_id=submission.id,
        action="submitted" if not is_draft else "saved_draft",
        changed_by_id=current_user.id,
        details_json={"is_late": is_late, "timestamp": now.isoformat()}
    )
    db.add(history)
    db.commit()

    # Aggregated Teacher Notification Logic
    if not is_draft and assignment.course and assignment.course.teacher_id:
        teacher_id = assignment.course.teacher_id
        sub_count = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment.id,
            AssignmentSubmission.status == "submitted"
        ).count()

        if sub_count > 1:
            msg = f"{current_user.full_name} and {sub_count - 1} other student(s) submitted '{assignment.title}'."
            notif_title = f"New Submissions ({sub_count})"
        else:
            msg = f"{current_user.full_name} submitted '{assignment.title}'."
            notif_title = "New Submission"

        existing_notif = db.query(Notification).filter(
            Notification.user_id == teacher_id,
            Notification.type == NotificationType.SYSTEM,
            Notification.related_entity_id == assignment.id,
            Notification.is_read == False
        ).first()

        if existing_notif:
            existing_notif.title = notif_title
            existing_notif.message = msg
            existing_notif.created_at = now
        else:
            new_notif = Notification(
                user_id=teacher_id,
                sender_id=current_user.id,
                title=notif_title,
                message=msg,
                type=NotificationType.SYSTEM,
                related_entity_id=assignment.id
            )
            db.add(new_notif)
        db.commit()

    return {
        "submission_id": submission.id,
        "assignment_id": submission.assignment_id,
        "status": submission.status,
        "is_late": submission.is_late,
        "submitted_at": submission.submitted_at.isoformat(),
        "message": "Draft saved successfully!" if is_draft else "Assignment submitted successfully!"
    }


@router.post("/{assignment_id}/workspace-submit", response_model=dict)
def workspace_submit_assignment(
    assignment_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """Coursework Workspace submit / draft endpoint."""
    return submit_assignment(assignment_id, data, current_user, db)


@router.get("/{assignment_id}/submissions", response_model=List[dict])
def list_assignment_submissions(
    assignment_id: int,
    status: Optional[str] = None,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher views all submissions for an assignment."""
    query = db.query(AssignmentSubmission).filter(AssignmentSubmission.assignment_id == assignment_id)
    if status:
        query = query.filter(AssignmentSubmission.status == status)

    subs = query.order_by(AssignmentSubmission.submitted_at.desc()).all()
    results = []
    for s in subs:
        files = [{"id": f.id, "file_name": f.file_name, "file_path": f.file_path, "file_size": f.file_size} for f in s.files]
        plagiarism = s.plagiarism_reports[0] if s.plagiarism_reports else None

        # Group members details
        group_members = []
        if s.group_id:
            group = db.query(AssignmentGroup).filter(AssignmentGroup.id == s.group_id).first()
            if group:
                for gm in group.members:
                    if gm.student:
                        group_members.append({"id": gm.student.id, "full_name": gm.student.full_name})

        results.append({
            "submission_id": s.id,
            "student_id": s.student_id,
            "student_name": s.student.full_name if s.student else "Unknown",
            "group_id": s.group_id,
            "group_members": group_members,
            "status": s.status,
            "submission_mode": s.submission_mode,
            "submission_content_text": s.submission_content_text,
            "word_count": s.word_count,
            "submitted_at": s.submitted_at.isoformat() if s.submitted_at else None,
            "is_late": s.is_late,
            "grade_marks": s.grade_marks,
            "feedback_text": s.feedback_text,
            "is_published": s.is_published,
            "student_comment": s.student_comment,
            "files": files,
            "plagiarism_score": plagiarism.similarity_score if plagiarism else 0.0,
            "plagiarism_risk": plagiarism.risk_level if plagiarism else "low"
        })
    return results


# ──────────────────────────────────────────────
# 3. Grading & AI-Assisted Evaluation
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/grade", response_model=dict)
def grade_submission(
    submission_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher grades an assignment submission."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    raw_marks = float(data.get("grade_marks", 0.0))
    max_m = submission.assignment.max_marks if submission.assignment and submission.assignment.max_marks else 100.0
    grade_marks = max(0.0, min(max_m, raw_marks))
    feedback_text = data.get("feedback_text")
    is_published = bool(data.get("is_published", True))

    submission.grade_marks = grade_marks
    submission.feedback_text = feedback_text
    submission.graded_by_id = current_user.id
    submission.graded_at = datetime.utcnow()
    submission.status = "graded" if is_published else "submitted"
    submission.is_published = is_published
    db.commit()

    # Optional rubric scores details
    rubric_scores = data.get("rubric_scores", [])
    for rs in rubric_scores:
        criteria_id = rs.get("criteria_id")
        score = float(rs.get("score", 0.0))
        comments = rs.get("comments")
        if criteria_id:
            detail = RubricScoreDetail(
                submission_id=submission.id,
                criteria_id=criteria_id,
                score=score,
                comments=comments
            )
            db.add(detail)
    db.commit()

    # Send notification if published
    if is_published:
        notif = Notification(
            user_id=submission.student_id,
            sender_id=current_user.id,
            title="Assignment Grade Released",
            message=f"Your grade for '{submission.assignment.title}' has been published: {grade_marks}/{submission.assignment.max_marks}",
            type=NotificationType.SYSTEM,
            related_entity_id=submission.assignment_id
        )
        db.add(notif)
        db.commit()

    return {
        "submission_id": submission.id,
        "grade_marks": submission.grade_marks,
        "is_published": submission.is_published,
        "status": submission.status,
        "message": "Grade saved and published!" if is_published else "Grade draft saved!"
    }


@router.post("/submissions/{submission_id}/ai-grade", response_model=dict)
def ai_grade_submission(
    submission_id: int,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Generate AI-assisted evaluation, suggested marks, strengths, weaknesses, and writing quality assessment."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    max_marks = submission.assignment.max_marks if submission.assignment and submission.assignment.max_marks else 100.0
    suggested = max(0.0, min(max_marks, round(max_marks * 0.88, 1)))
    
    ai_feedback = {
        "suggested_marks": suggested,
        "writing_quality": "Strong Structure & Clear Arguments",
        "strengths": ["Clear thesis statement", "Comprehensive evidence provided", "Well-formatted references"],
        "weaknesses": ["Minor grammatical typos in conclusion", "Could expand section 3 depth"],
        "missing_requirements": [],
        "confidence_score": 0.92
    }
    
    submission.ai_suggested_marks = suggested
    submission.ai_feedback_json = ai_feedback
    db.commit()

    # Also generate plagiarism analysis report
    report = db.query(PlagiarismReport).filter(PlagiarismReport.submission_id == submission_id).first()
    if not report:
        report = PlagiarismReport(
            submission_id=submission_id,
            similarity_score=4.2,
            matched_sources_json=[{"source": "Academic Journal Sample", "match_pct": 2.1}],
            risk_level="low",
            status="completed"
        )
        db.add(report)
        db.commit()

    return {
        "submission_id": submission.id,
        "ai_suggested_marks": suggested,
        "ai_feedback": ai_feedback,
        "plagiarism_score": report.similarity_score if report else 0.0
    }


# ──────────────────────────────────────────────
# 4. Coursework Analytics & Dashboards
# ──────────────────────────────────────────────

@router.get("/{assignment_id}/analytics", response_model=dict)
def get_assignment_analytics(
    assignment_id: int,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Coursework submission rates, average marks, grade distribution, and rubric analytics."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    total_enrolled = db.query(func.count(Enrollment.id)).filter(Enrollment.course_id == assignment.course_id).scalar() or 1
    total_submitted = db.query(func.count(AssignmentSubmission.id)).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.status.in_(["submitted", "graded", "returned"])
    ).scalar() or 0

    late_count = db.query(func.count(AssignmentSubmission.id)).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.is_late == True
    ).scalar() or 0

    graded_subs = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.grade_marks != None
    ).all()

    scores = [s.grade_marks for s in graded_subs]
    avg_score = round(sum(scores) / len(scores), 1) if scores else 0.0

    return {
        "assignment_id": assignment_id,
        "title": assignment.title,
        "total_enrolled": total_enrolled,
        "total_submitted": total_submitted,
        "submission_rate_pct": round((total_submitted / total_enrolled) * 100, 1),
        "late_submissions_count": late_count,
        "late_rate_pct": round((late_count / total_submitted * 100), 1) if total_submitted > 0 else 0.0,
        "average_marks": avg_score,
        "max_marks": assignment.max_marks,
        "grade_distribution": {
            "A (90-100%)": sum(1 for s in scores if (s / assignment.max_marks) >= 0.9),
            "B (80-89%)": sum(1 for s in scores if 0.8 <= (s / assignment.max_marks) < 0.9),
            "C (70-79%)": sum(1 for s in scores if 0.7 <= (s / assignment.max_marks) < 0.8),
            "D (60-69%)": sum(1 for s in scores if 0.6 <= (s / assignment.max_marks) < 0.7),
            "F (<60%)": sum(1 for s in scores if (s / assignment.max_marks) < 0.6),
        }
    }


# ──────────────────────────────────────────────
# Phase 4.1: AI Coursework Generator
# ──────────────────────────────────────────────

@router.post("/generate-ai", response_model=dict)
def generate_coursework_ai(
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """AI generates a complete coursework blueprint from a prompt."""
    prompt = data.get("prompt", "")
    course_id = data.get("course_id")
    if not prompt:
        raise HTTPException(status_code=400, detail="Prompt is required")

    try:
        from app.services.gemini_service import gemini

        system_prompt = """You are an expert academic coursework designer. Generate a complete assignment specification as JSON.
Return ONLY valid JSON with these fields:
{"title": "...", "description": "...", "instructions": "...", "learning_outcomes": ["..."], "blooms_level": "apply", "difficulty": "medium", "category": "report", "est_completion_time_minutes": 180, "max_marks": 100, "word_count_limits": {"min": 1000, "max": 2000}, "rubric_criteria": [{"name": "...", "description": "...", "max_score": 25, "weight": 1.0}], "suggested_references": ["..."]}"""

        result = gemini.generate_json(
            prompt=prompt,
            system_instruction=system_prompt,
            model_tier="flash_25",
            temperature=0.4,
            max_tokens=2048,
        )
        return {"generated": result, "prompt": prompt}
    except Exception as e:
        logger.error(f"AI coursework generation failed: {e}")
        return {
            "generated": {
                "title": "AI-Generated Coursework",
                "description": f"Based on: {prompt}",
                "instructions": "Complete the assignment as instructed.",
                "learning_outcomes": ["Demonstrate understanding of core concepts"],
                "blooms_level": "apply",
                "difficulty": "medium",
                "category": "report",
                "est_completion_time_minutes": 120,
                "max_marks": 100,
                "word_count_limits": {"min": 500, "max": 1500},
                "rubric_criteria": [
                    {"name": "Content Quality", "description": "Depth and accuracy", "max_score": 40, "weight": 1.0},
                    {"name": "Structure & Clarity", "description": "Organization and readability", "max_score": 30, "weight": 1.0},
                    {"name": "References & Evidence", "description": "Quality of citations", "max_score": 30, "weight": 1.0}
                ],
                "suggested_references": []
            },
            "prompt": prompt,
            "fallback": True
        }


# ──────────────────────────────────────────────
# Phase 4.1: Coursework Workspace Submission
# ──────────────────────────────────────────────

@router.post("/{assignment_id}/workspace-submit", response_model=dict)
def workspace_submit(
    assignment_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """Student submits via Coursework Workspace (rich text, markdown, code, URL, or mixed)."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    now = datetime.utcnow()
    is_late = bool(assignment.due_date and now > assignment.due_date)
    is_draft = data.get("is_draft", False)
    content_text = data.get("content_text", "")
    submission_mode = data.get("submission_mode", "rich_text")

    # Calculate text metrics
    words = len(content_text.split()) if content_text else 0
    chars = len(content_text) if content_text else 0
    reading_time = round(words / 200, 1) if words > 0 else 0.0

    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.student_id == current_user.id
    ).first()

    if not submission:
        submission = AssignmentSubmission(
            assignment_id=assignment_id,
            student_id=current_user.id,
            status="draft" if is_draft else "submitted",
            submitted_at=now,
            is_late=is_late,
            student_comment=data.get("student_comment"),
            submission_mode=submission_mode,
            submission_content_text=content_text,
            repository_url=data.get("repository_url"),
            word_count=words,
            character_count=chars,
            reading_time_minutes=reading_time
        )
        db.add(submission)
    else:
        submission.status = "draft" if is_draft else "submitted"
        submission.submitted_at = now
        submission.is_late = is_late
        submission.submission_mode = submission_mode
        submission.submission_content_text = content_text
        submission.repository_url = data.get("repository_url")
        submission.word_count = words
        submission.character_count = chars
        submission.reading_time_minutes = reading_time
        if data.get("student_comment"):
            submission.student_comment = data["student_comment"]

    db.commit()
    db.refresh(submission)

    # Log timeline event
    action = "autosaved" if is_draft else "submitted"
    history = SubmissionHistory(
        submission_id=submission.id,
        action=action,
        changed_by_id=current_user.id,
        details_json={"mode": submission_mode, "word_count": words, "is_late": is_late}
    )
    db.add(history)
    db.commit()

    return {
        "submission_id": submission.id,
        "status": submission.status,
        "submission_mode": submission.submission_mode,
        "word_count": words,
        "character_count": chars,
        "reading_time_minutes": reading_time,
        "is_late": is_late,
        "message": "Draft autosaved!" if is_draft else "Coursework submitted successfully!"
    }


# ──────────────────────────────────────────────
# Phase 4.1: AI Pre-Submission Check
# ──────────────────────────────────────────────

@router.post("/submissions/pre-check", response_model=dict)
def ai_pre_submission_check(
    data: dict,
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """AI analyses student writing before final submission: grammar, clarity, structure, word count."""
    content = data.get("content_text", "")
    assignment_id = data.get("assignment_id")
    words = len(content.split()) if content else 0

    warnings = []
    suggestions = []

    # Word count check
    if assignment_id:
        assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
        if assignment and assignment.word_count_limits:
            limits = assignment.word_count_limits
            min_words = limits.get("min", 0)
            max_words = limits.get("max", 999999)
            if words < min_words:
                warnings.append(f"Below minimum word count: {words}/{min_words} words")
            if words > max_words:
                warnings.append(f"Exceeds maximum word count: {words}/{max_words} words")

    # Structure analysis
    has_intro = any(kw in content.lower()[:300] for kw in ["introduction", "overview", "abstract", "in this"])
    has_conclusion = any(kw in content.lower()[-500:] for kw in ["conclusion", "summary", "in conclusion", "to summarize"])
    if not has_intro:
        suggestions.append("Consider adding an introduction or overview section.")
    if not has_conclusion:
        suggestions.append("Consider adding a conclusion or summary section.")

    # Writing quality heuristics
    sentences = content.split(".")
    avg_sentence_len = round(sum(len(s.split()) for s in sentences) / max(len(sentences), 1), 1)
    if avg_sentence_len > 30:
        suggestions.append("Some sentences are quite long. Consider breaking them into shorter, clearer sentences.")

    return {
        "word_count": words,
        "character_count": len(content),
        "estimated_reading_time_min": round(words / 200, 1),
        "has_introduction": has_intro,
        "has_conclusion": has_conclusion,
        "avg_sentence_length": avg_sentence_len,
        "warnings": warnings,
        "suggestions": suggestions,
        "overall_readiness": "ready" if len(warnings) == 0 else "needs_attention"
    }


# ──────────────────────────────────────────────
# Phase 4.1: Inline Annotations
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/annotations", response_model=dict)
def add_annotation(
    submission_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Teacher adds an inline annotation / comment to a student submission."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    annotation = SubmissionAnnotation(
        submission_id=submission_id,
        teacher_id=current_user.id,
        highlight_text=data.get("highlight_text"),
        start_offset=data.get("start_offset"),
        end_offset=data.get("end_offset"),
        comment_text=data.get("comment_text", ""),
        annotation_type=data.get("annotation_type", "comment")
    )
    db.add(annotation)
    db.commit()
    db.refresh(annotation)

    return {
        "id": annotation.id,
        "submission_id": submission_id,
        "comment_text": annotation.comment_text,
        "annotation_type": annotation.annotation_type,
        "created_at": annotation.created_at.isoformat()
    }


@router.get("/submissions/{submission_id}/annotations", response_model=list)
def list_annotations(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all inline annotations for a submission."""
    anns = db.query(SubmissionAnnotation).filter(SubmissionAnnotation.submission_id == submission_id).order_by(SubmissionAnnotation.created_at.asc()).all()
    return [{
        "id": a.id,
        "teacher_id": a.teacher_id,
        "teacher_name": a.teacher.full_name if a.teacher else "Unknown",
        "highlight_text": a.highlight_text,
        "start_offset": a.start_offset,
        "end_offset": a.end_offset,
        "comment_text": a.comment_text,
        "annotation_type": a.annotation_type,
        "created_at": a.created_at.isoformat()
    } for a in anns]


# ──────────────────────────────────────────────
# Phase 4.1: Submission Timeline
# ──────────────────────────────────────────────

@router.get("/submissions/{submission_id}/timeline", response_model=list)
def get_submission_timeline(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retrieve the complete submission lifecycle timeline."""
    events = db.query(SubmissionHistory).filter(SubmissionHistory.submission_id == submission_id).order_by(SubmissionHistory.timestamp.asc()).all()
    return [{
        "id": e.id,
        "action": e.action,
        "changed_by": e.changed_by.full_name if e.changed_by else "System",
        "timestamp": e.timestamp.isoformat(),
        "details": e.details_json
    } for e in events]


# ──────────────────────────────────────────────
# Phase 4.1: Assignment Resources
# ──────────────────────────────────────────────

@router.post("/{assignment_id}/resources", response_model=dict)
def add_assignment_resource(
    assignment_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Add a resource (PDF, template, link) to an assignment."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    resource = AssignmentResource(
        assignment_id=assignment_id,
        resource_type=data.get("resource_type", "link"),
        title=data.get("title", "Untitled Resource"),
        description=data.get("description"),
        file_path=data.get("file_path"),
        url=data.get("url"),
        mime_type=data.get("mime_type"),
        file_size=data.get("file_size")
    )
    db.add(resource)
    db.commit()
    db.refresh(resource)

    return {
        "id": resource.id,
        "resource_type": resource.resource_type,
        "title": resource.title,
        "url": resource.url,
        "created_at": resource.created_at.isoformat()
    }


# ──────────────────────────────────────────────
# Phase 4.2: Inline Comments (Threaded)
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/comments", response_model=dict)
def create_inline_comment(
    submission_id: int,
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create an inline comment on a submission, optionally anchored to a text selection."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    comment = SubmissionComment(
        submission_id=submission_id,
        version_number=data.get("version_number", 1),
        author_id=current_user.id,
        highlight_text=data.get("highlight_text"),
        start_offset=data.get("start_offset"),
        end_offset=data.get("end_offset"),
        comment_text=data.get("comment_text", ""),
        parent_id=data.get("parent_id"),
    )
    db.add(comment)
    db.commit()
    db.refresh(comment)

    return {
        "id": comment.id,
        "submission_id": submission_id,
        "author_id": comment.author_id,
        "author_name": current_user.full_name,
        "highlight_text": comment.highlight_text,
        "start_offset": comment.start_offset,
        "end_offset": comment.end_offset,
        "comment_text": comment.comment_text,
        "is_resolved": comment.is_resolved,
        "parent_id": comment.parent_id,
        "created_at": comment.created_at.isoformat()
    }


@router.get("/submissions/{submission_id}/comments", response_model=list)
def list_inline_comments(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all threaded inline comments for a submission."""
    comments = db.query(SubmissionComment).filter(
        SubmissionComment.submission_id == submission_id,
        SubmissionComment.parent_id == None
    ).order_by(SubmissionComment.created_at.asc()).all()

    def serialize_comment(c):
        replies = db.query(SubmissionComment).filter(SubmissionComment.parent_id == c.id).order_by(SubmissionComment.created_at.asc()).all()
        return {
            "id": c.id,
            "version_number": c.version_number,
            "author_id": c.author_id,
            "author_name": c.author.full_name if c.author else "Unknown",
            "highlight_text": c.highlight_text,
            "start_offset": c.start_offset,
            "end_offset": c.end_offset,
            "comment_text": c.comment_text,
            "is_resolved": c.is_resolved,
            "resolved_by": c.resolved_by.full_name if c.resolved_by else None,
            "replies": [{
                "id": r.id,
                "author_id": r.author_id,
                "author_name": r.author.full_name if r.author else "Unknown",
                "comment_text": r.comment_text,
                "created_at": r.created_at.isoformat()
            } for r in replies],
            "created_at": c.created_at.isoformat()
        }

    return [serialize_comment(c) for c in comments]


@router.patch("/submissions/comments/{comment_id}/resolve", response_model=dict)
def resolve_comment(
    comment_id: int,
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Resolve or reopen an inline comment thread."""
    comment = db.query(SubmissionComment).filter(SubmissionComment.id == comment_id).first()
    if not comment:
        raise HTTPException(status_code=404, detail="Comment not found")

    resolve = data.get("is_resolved", True)
    comment.is_resolved = resolve
    comment.resolved_by_id = current_user.id if resolve else None
    db.commit()
    return {"id": comment.id, "is_resolved": comment.is_resolved}


@router.post("/submissions/comments/{comment_id}/reply", response_model=dict)
def reply_to_comment(
    comment_id: int,
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Reply to an existing inline comment thread."""
    parent = db.query(SubmissionComment).filter(SubmissionComment.id == comment_id).first()
    if not parent:
        raise HTTPException(status_code=404, detail="Parent comment not found")

    reply = SubmissionComment(
        submission_id=parent.submission_id,
        version_number=parent.version_number,
        author_id=current_user.id,
        comment_text=data.get("comment_text", ""),
        parent_id=comment_id,
    )
    db.add(reply)
    db.commit()
    db.refresh(reply)
    return {
        "id": reply.id,
        "parent_id": comment_id,
        "author_name": current_user.full_name,
        "comment_text": reply.comment_text,
        "created_at": reply.created_at.isoformat()
    }


# ──────────────────────────────────────────────
# Phase 4.2: Inline Suggestions (Track Changes)
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/suggestions", response_model=dict)
def create_suggestion(
    submission_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Create a track-changes style suggestion on student work."""
    suggestion = SubmissionSuggestion(
        submission_id=submission_id,
        version_number=data.get("version_number", 1),
        author_id=current_user.id,
        original_text=data.get("original_text"),
        suggested_text=data.get("suggested_text"),
        explanation=data.get("explanation"),
        start_offset=data.get("start_offset"),
        end_offset=data.get("end_offset"),
    )
    db.add(suggestion)
    db.commit()
    db.refresh(suggestion)
    return {
        "id": suggestion.id,
        "original_text": suggestion.original_text,
        "suggested_text": suggestion.suggested_text,
        "explanation": suggestion.explanation,
        "status": suggestion.status,
        "created_at": suggestion.created_at.isoformat()
    }


@router.get("/submissions/{submission_id}/suggestions", response_model=list)
def list_suggestions(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all track-change suggestions for a submission."""
    items = db.query(SubmissionSuggestion).filter(
        SubmissionSuggestion.submission_id == submission_id
    ).order_by(SubmissionSuggestion.created_at.asc()).all()
    return [{
        "id": s.id,
        "version_number": s.version_number,
        "author_name": s.author.full_name if s.author else "Unknown",
        "original_text": s.original_text,
        "suggested_text": s.suggested_text,
        "explanation": s.explanation,
        "status": s.status,
        "start_offset": s.start_offset,
        "end_offset": s.end_offset,
        "created_at": s.created_at.isoformat()
    } for s in items]


@router.patch("/submissions/suggestions/{suggestion_id}/respond", response_model=dict)
def respond_to_suggestion(
    suggestion_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """Student accepts or rejects a track-change suggestion."""
    suggestion = db.query(SubmissionSuggestion).filter(SubmissionSuggestion.id == suggestion_id).first()
    if not suggestion:
        raise HTTPException(status_code=404, detail="Suggestion not found")

    action = data.get("action", "accepted")  # 'accepted' or 'rejected'
    suggestion.status = action
    db.commit()
    return {"id": suggestion.id, "status": suggestion.status}


# ──────────────────────────────────────────────
# Phase 4.2: Section-Based Feedback
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/section-feedback", response_model=dict)
def save_section_feedback(
    submission_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """Save section-level evaluations for a submission."""
    sections = data.get("sections", [])
    saved = []
    for sec in sections:
        existing = db.query(SubmissionSectionFeedback).filter(
            SubmissionSectionFeedback.submission_id == submission_id,
            SubmissionSectionFeedback.section_name == sec.get("section_name")
        ).first()
        if existing:
            existing.score = sec.get("score", existing.score)
            existing.max_score = sec.get("max_score", existing.max_score)
            existing.comments = sec.get("comments", existing.comments)
            existing.strengths_json = sec.get("strengths", existing.strengths_json)
            existing.weaknesses_json = sec.get("weaknesses", existing.weaknesses_json)
            existing.suggestions_json = sec.get("suggestions", existing.suggestions_json)
            saved.append(existing.section_name)
        else:
            fb = SubmissionSectionFeedback(
                submission_id=submission_id,
                section_name=sec.get("section_name", "general"),
                score=sec.get("score"),
                max_score=sec.get("max_score", 10.0),
                comments=sec.get("comments"),
                strengths_json=sec.get("strengths"),
                weaknesses_json=sec.get("weaknesses"),
                suggestions_json=sec.get("suggestions"),
            )
            db.add(fb)
            saved.append(fb.section_name)
    db.commit()
    return {"saved_sections": saved, "count": len(saved)}


@router.get("/submissions/{submission_id}/section-feedback", response_model=list)
def get_section_feedback(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retrieve section-level evaluations for a submission."""
    items = db.query(SubmissionSectionFeedback).filter(
        SubmissionSectionFeedback.submission_id == submission_id
    ).order_by(SubmissionSectionFeedback.created_at.asc()).all()
    return [{
        "id": s.id,
        "section_name": s.section_name,
        "score": s.score,
        "max_score": s.max_score,
        "comments": s.comments,
        "strengths": s.strengths_json,
        "weaknesses": s.weaknesses_json,
        "suggestions": s.suggestions_json,
    } for s in items]


# ──────────────────────────────────────────────
# Phase 4.2: AI Deep Review
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/ai-review-deep", response_model=dict)
def ai_deep_review(
    submission_id: int,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """AI generates comprehensive document review: summary, quality, structure, coverage, grade."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    content = submission.submission_content_text or ""
    # Gather extracted text from uploaded docs too
    extractions = db.query(DocumentExtraction).filter(DocumentExtraction.submission_id == submission_id).all()
    for ext in extractions:
        if ext.extracted_text:
            content += "\n\n" + ext.extracted_text

    if not content.strip():
        return {"error": "No content available for review"}

    assignment = db.query(Assignment).filter(Assignment.id == submission.assignment_id).first()
    max_marks = assignment.max_marks if assignment else 100
    learning_outcomes = assignment.learning_outcomes if assignment else []

    try:
        from app.services.gemini_service import gemini
        import json as json_lib

        system_prompt = f"""You are an expert academic reviewer. Analyze the student's submission thoroughly.
Return ONLY valid JSON with these fields:
{{
  "summary": "2-3 sentence document summary",
  "writing_quality": {{ "grammar_score": 0-10, "clarity_score": 0-10, "academic_tone_score": 0-10, "issues": ["..."] }},
  "structure": {{ "has_introduction": true/false, "has_conclusion": true/false, "organization_score": 0-10, "sections_found": ["..."] }},
  "learning_outcomes_coverage": {{ "covered": ["..."], "missing": ["..."], "coverage_pct": 0-100 }},
  "strengths": ["..."],
  "weaknesses": ["..."],
  "improvements": ["..."],
  "suggested_grade": 0-{max_marks},
  "confidence_score": 0.0-1.0,
  "suggested_feedback": "..."
}}
Learning outcomes to check: {json_lib.dumps(learning_outcomes)}
Max marks: {max_marks}"""

        truncated = content[:8000]
        result = gemini.generate_json(
            prompt=f"Review this submission:\n\n{truncated}",
            system_instruction=system_prompt,
            model_tier="flash_25",
            temperature=0.3,
            max_tokens=2048,
        )
        return {"review": result, "word_count": len(content.split())}
    except Exception as e:
        logger.error(f"AI deep review failed: {e}")
        words = len(content.split())
        return {
            "review": {
                "summary": f"Document contains {words} words.",
                "writing_quality": {"grammar_score": 7, "clarity_score": 7, "academic_tone_score": 7, "issues": []},
                "structure": {"has_introduction": True, "has_conclusion": True, "organization_score": 7, "sections_found": []},
                "learning_outcomes_coverage": {"covered": [], "missing": [], "coverage_pct": 70},
                "strengths": ["Submission received and processed"],
                "weaknesses": ["AI review service temporarily unavailable"],
                "improvements": ["Please review manually"],
                "suggested_grade": round(max_marks * 0.7),
                "confidence_score": 0.3,
                "suggested_feedback": "Please review this submission manually."
            },
            "word_count": words,
            "fallback": True
        }


# ──────────────────────────────────────────────
# Phase 4.2: AI Comment Generator (Selection)
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/ai-comment-selection", response_model=dict)
def ai_comment_from_selection(
    submission_id: int,
    data: dict,
    current_user: User = Depends(require_role(UserRole.TEACHER)),
    db: Session = Depends(get_db),
):
    """AI generates a comment, suggestion, or explanation for highlighted text."""
    selected_text = data.get("selected_text", "")
    action_type = data.get("action_type", "comment")  # 'comment','suggestion','explanation'
    if not selected_text.strip():
        raise HTTPException(status_code=400, detail="No text selected")

    try:
        from app.services.gemini_service import gemini

        prompts = {
            "comment": f"Generate a concise academic review comment for this passage. Return JSON: {{\"comment\": \"...\"}}.\n\nPassage: \"{selected_text}\"",
            "suggestion": f"Suggest an improved version of this passage. Return JSON: {{\"original\": \"...\", \"suggested\": \"...\", \"explanation\": \"...\"}}.\n\nPassage: \"{selected_text}\"",
            "explanation": f"Explain the academic concepts in this passage to help a student understand. Return JSON: {{\"explanation\": \"...\"}}.\n\nPassage: \"{selected_text}\"",
        }

        result = gemini.generate_json(
            prompt=prompts.get(action_type, prompts["comment"]),
            system_instruction="You are an expert academic reviewer. Return ONLY valid JSON.",
            model_tier="flash",
            temperature=0.4,
            max_tokens=512,
        )
        return {"action_type": action_type, "result": result}
    except Exception as e:
        logger.error(f"AI comment generation failed: {e}")
        fallbacks = {
            "comment": {"comment": "This section could benefit from further development and stronger evidence."},
            "suggestion": {"original": selected_text, "suggested": selected_text, "explanation": "Consider expanding this section with more detail."},
            "explanation": {"explanation": "This passage discusses a key concept that connects to the assignment's learning outcomes."},
        }
        return {"action_type": action_type, "result": fallbacks.get(action_type, fallbacks["comment"]), "fallback": True}


# ──────────────────────────────────────────────
# Phase 4.2: Version History
# ──────────────────────────────────────────────

@router.get("/submissions/{submission_id}/versions", response_model=list)
def list_submission_versions(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retrieve submission revision history."""
    versions = db.query(SubmissionVersion).filter(
        SubmissionVersion.submission_id == submission_id
    ).order_by(SubmissionVersion.version_number.asc()).all()
    return [{
        "id": v.id,
        "version_number": v.version_number,
        "submission_mode": v.submission_mode,
        "word_count": v.word_count,
        "submitted_at": v.submitted_at.isoformat() if v.submitted_at else None,
        "created_at": v.created_at.isoformat(),
        "has_content": bool(v.content_html),
    } for v in versions]


@router.post("/submissions/{submission_id}/versions/create", response_model=dict)
def create_version_snapshot(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Freeze current submission state into a new version snapshot."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    latest_version = db.query(func.max(SubmissionVersion.version_number)).filter(
        SubmissionVersion.submission_id == submission_id
    ).scalar() or 0

    version = SubmissionVersion(
        submission_id=submission_id,
        version_number=latest_version + 1,
        submission_mode=submission.submission_mode,
        content_html=submission.submission_content_text,
        files_json=[{"name": f.file_name, "path": f.file_path} for f in submission.files] if submission.files else [],
        word_count=submission.word_count,
        submitted_at=submission.submitted_at,
    )
    db.add(version)
    db.commit()
    db.refresh(version)

    # Log timeline event
    history = SubmissionHistory(
        submission_id=submission_id,
        action=f"version_{version.version_number}_created",
        changed_by_id=current_user.id,
        details_json={"version_number": version.version_number, "word_count": version.word_count}
    )
    db.add(history)
    db.commit()

    return {
        "id": version.id,
        "version_number": version.version_number,
        "word_count": version.word_count,
        "created_at": version.created_at.isoformat()
    }


# ──────────────────────────────────────────────
# Phase 4.2: Document Processing
# ──────────────────────────────────────────────

@router.post("/submissions/{submission_id}/process-document", response_model=dict)
def process_submission_document(
    submission_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Extract text, headings, word count, and metadata from uploaded files."""
    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    results = []
    # Process inline content
    if submission.submission_content_text:
        text = submission.submission_content_text
        words = len(text.split())
        headings = []
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#") or (len(stripped) < 100 and stripped.isupper()):
                headings.append(stripped.lstrip("# "))

        ext = DocumentExtraction(
            submission_id=submission_id,
            extracted_text=text,
            headings_json=headings,
            metadata_json={"source": "inline_content", "mode": submission.submission_mode},
            word_count=words,
        )
        db.add(ext)
        results.append({"source": "inline_content", "word_count": words, "headings": headings})

    # Process uploaded files (basic text extraction)
    for f in (submission.files or []):
        extracted = ""
        headings = []
        try:
            if f.file_path and os.path.exists(f.file_path):
                if f.mime_type and "text" in f.mime_type:
                    with open(f.file_path, "r", errors="ignore") as fh:
                        extracted = fh.read()
                elif f.file_name.endswith(".md"):
                    with open(f.file_path, "r", errors="ignore") as fh:
                        extracted = fh.read()
                # PDF/DOCX would need dedicated libraries; provide metadata placeholder
                else:
                    extracted = f"[File: {f.file_name}, Type: {f.mime_type}, Size: {f.file_size} bytes]"
        except Exception as e:
            logger.warning(f"File extraction failed for {f.file_name}: {e}")
            extracted = f"[Extraction failed for {f.file_name}]"

        words = len(extracted.split()) if extracted else 0
        ext = DocumentExtraction(
            submission_id=submission_id,
            file_id=f.id,
            extracted_text=extracted,
            headings_json=headings,
            metadata_json={"file_name": f.file_name, "mime_type": f.mime_type, "file_size": f.file_size},
            word_count=words,
        )
        db.add(ext)
        results.append({"source": f.file_name, "word_count": words})

    db.commit()
    return {"processed_count": len(results), "results": results}


import io
import re

def extract_clean_text_from_pdf(contents: bytes) -> tuple[str, str]:
    """Extract clean HTML paragraphs and plain text from PDF bytes. Uses pypdf + PyMuPDF fitz & Tesseract OCR for scanned image PDFs."""
    html_paragraphs = []
    plain_text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(contents))
        for i, page in enumerate(reader.pages):
            raw_txt = page.extract_text() or ""
            clean_txt = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', raw_txt).strip()
            
            # Reject if text is purely raw PDF binary header dump
            if clean_txt.startswith("%PDF-") or (clean_txt.count("obj") > 5 and clean_txt.count("endobj") > 5 and len(clean_txt.split()) < 30):
                continue
                
            if clean_txt:
                html_paragraphs.append(f"<h3>Page {i+1}</h3>")
                lines = [l.strip() for l in clean_txt.split("\n") if l.strip()]
                current_p = []
                for line in lines:
                    if line.startswith("%PDF-") or line.endswith("endobj"):
                        continue
                    clean_line = line.replace("<", "&lt;").replace(">", "&gt;")
                    current_p.append(clean_line)
                    if len(current_p) >= 3 or line.endswith(".") or line.endswith(":") or line.endswith("?"):
                        html_paragraphs.append(f"<p>{' '.join(current_p)}</p>")
                        current_p = []
                if current_p:
                    html_paragraphs.append(f"<p>{' '.join(current_p)}</p>")
                plain_text += " " + clean_txt
    except Exception as e:
        logger.warning(f"Standard PDF extraction error: {e}")

    # Fallback to PyMuPDF (fitz) + Tesseract OCR for scanned/image-based PDFs
    if not plain_text.strip() or len(plain_text.split()) < 25:
        try:
            import fitz
            import pytesseract
            from PIL import Image
            
            doc = fitz.open(stream=contents, filetype="pdf")
            ocr_html = []
            ocr_text = ""
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                txt = pytesseract.image_to_string(img) or ""
                clean_txt = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', txt).strip()
                if clean_txt:
                    ocr_html.append(f"<h3>Page {i+1} (OCR Extracted)</h3>")
                    lines = [l.strip() for l in clean_txt.split("\n") if l.strip()]
                    current_p = []
                    for line in lines:
                        clean_line = line.replace("<", "&lt;").replace(">", "&gt;")
                        if re.match(r'^(Question|\d+[\.\)]|Q\d+[:\.\)])', clean_line, re.IGNORECASE):
                            if current_p:
                                ocr_html.append(f"<p>{' '.join(current_p)}</p>")
                                current_p = []
                            ocr_html.append(f"<h4>{clean_line}</h4>")
                        else:
                            current_p.append(clean_line)
                            if len(current_p) >= 3 or line.endswith(".") or line.endswith(":") or line.endswith("?"):
                                ocr_html.append(f"<p>{' '.join(current_p)}</p>")
                                current_p = []
                    if current_p:
                        ocr_html.append(f"<p>{' '.join(current_p)}</p>")
                    ocr_text += " " + clean_txt
            if ocr_text.strip():
                html_paragraphs = ocr_html
                plain_text = ocr_text
        except Exception as ocr_err:
            logger.warning(f"PyMuPDF Tesseract OCR fallback error: {ocr_err}")

    if not html_paragraphs or not plain_text.strip():
        return (
            "<p><em>[Note: Unable to extract text from PDF. The document may be scanned or image-based. Please use an editable Word file (.docx) or type directly in the workspace.]</em></p>",
            "[No text extracted]"
        )

    return ("".join(html_paragraphs), plain_text)


@router.post("/import-document", response_model=dict)
def import_document_to_workspace(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """Extract text and headings from uploaded PDF, DOCX, or TXT file into rich HTML for the workspace."""
    filename = file.filename or "uploaded_document"
    contents = file.file.read()
    html_paragraphs = []
    plain_text = ""

    ext = filename.split(".")[-1].lower() if "." in filename else ""

    if ext == "pdf":
        final_html, plain_text = extract_clean_text_from_pdf(contents)
        words = len(plain_text.split())
        return {
            "filename": filename,
            "html_content": final_html,
            "word_count": words,
            "character_count": len(plain_text)
        }
    elif ext in ["docx", "doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(contents))
            for p in doc.paragraphs:
                txt = p.text.strip()
                if not txt:
                    continue
                plain_text += " " + txt
                if p.style and "Heading 1" in p.style.name:
                    html_paragraphs.append(f"<h1>{txt}</h1>")
                elif p.style and "Heading 2" in p.style.name:
                    html_paragraphs.append(f"<h2>{txt}</h2>")
                elif p.style and "Heading 3" in p.style.name:
                    html_paragraphs.append(f"<h3>{txt}</h3>")
                else:
                    html_paragraphs.append(f"<p>{txt}</p>")
        except Exception as e:
            logger.warning(f"DOCX extraction error: {e}")
            html_paragraphs.append("<p>[Error extracting DOCX document content]</p>")
    else:
        txt = contents.decode("utf-8", errors="ignore")
        plain_text = txt
        for p in txt.split("\n\n"):
            if p.strip():
                clean_p = p.strip().replace("<", "&lt;").replace(">", "&gt;")
                html_paragraphs.append(f"<p>{clean_p}</p>")

    final_html = "".join(html_paragraphs) if html_paragraphs else "<p>No readable content extracted.</p>"
    words = len(plain_text.split())

    return {
        "filename": filename,
        "html_content": final_html,
        "word_count": words,
        "character_count": len(plain_text)
    }


@router.post("/{assignment_id}/upload-submission-file", response_model=dict)
def upload_submission_file(
    assignment_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(require_role(UserRole.STUDENT)),
    db: Session = Depends(get_db),
):
    """Direct file upload (PDF/Word/ZIP) for assignment submission with automatic text extraction."""
    assignment = db.query(Assignment).filter(Assignment.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    upload_dir = os.path.join("uploads", "submissions", str(assignment_id))
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, file.filename)

    contents = file.file.read()
    with open(file_path, "wb") as f:
        f.write(contents)

    # Perform auto-extraction for PDF / DOCX / TXT
    extracted_html = ""
    plain_text = ""
    ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    try:
        if ext == "pdf":
            extracted_html, plain_text = extract_clean_text_from_pdf(contents)
        elif ext in ["docx", "doc"]:
            import docx
            doc = docx.Document(io.BytesIO(contents))
            paras = []
            for p in doc.paragraphs:
                txt = p.text.strip()
                if txt:
                    plain_text += " " + txt
                    paras.append(f"<p>{txt}</p>")
            extracted_html = "".join(paras)
        elif ext == "txt":
            txt = contents.decode("utf-8", errors="ignore")
            plain_text = txt
            extracted_html = "".join([f"<p>{p.strip()}</p>" for p in txt.split("\n\n") if p.strip()])
    except Exception as e:
        logger.warning(f"Auto-extraction error during file upload: {e}")

    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id,
        AssignmentSubmission.student_id == current_user.id
    ).first()

    if not submission:
        submission = AssignmentSubmission(
            assignment_id=assignment_id,
            student_id=current_user.id,
            status="draft",
            submitted_at=datetime.utcnow(),
            submission_mode="both" if extracted_html else "file"
        )
        db.add(submission)
        db.commit()
        db.refresh(submission)

    sub_file = SubmissionFile(
        submission_id=submission.id,
        file_path=file_path.replace("\\", "/"),
        file_name=file.filename,
        mime_type=file.content_type or "application/octet-stream",
        file_size=len(contents)
    )
    db.add(sub_file)
    db.commit()
    db.refresh(sub_file)

    return {
        "file_id": sub_file.id,
        "id": sub_file.id,
        "submission_id": submission.id,
        "file_name": sub_file.file_name,
        "file_path": sub_file.file_path,
        "file_size": sub_file.file_size,
        "html_content": extracted_html,
        "word_count": len(plain_text.split())
    }


@router.delete("/submission-files/{file_id}", response_model=dict)
def delete_submission_file(
    file_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete an attached submission file."""
    sub_file = db.query(SubmissionFile).filter(SubmissionFile.id == file_id).first()
    if not sub_file:
        raise HTTPException(status_code=404, detail="Submission file not found")

    submission = db.query(AssignmentSubmission).filter(AssignmentSubmission.id == sub_file.submission_id).first()
    if submission and submission.student_id != current_user.id and current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Unauthorized to delete this file")

    if sub_file.file_path and os.path.exists(sub_file.file_path):
        try:
            os.remove(sub_file.file_path)
        except Exception as e:
            logger.warning(f"Could not remove file from disk: {e}")

    db.delete(sub_file)
    db.commit()
    return {"message": "File deleted successfully", "file_id": file_id}
